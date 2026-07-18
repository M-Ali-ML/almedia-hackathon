"""Ingest a GDPdU-style dossier ZIP into one DuckDB database per batch.

Provenance rules (see docs/mvp.md):
- every source file gets a stable ``document_id`` (doc-001, doc-002, ... in
  sorted path order) recorded in the ``documents`` table
- every tabular row carries ``_row_id`` = physical line number in the source
  file (GDPdU txt files have no header, so data starts at line 1; CSVs start
  at line 2; XLSX rows use the spreadsheet row number)
- prose documents (DOCX/PDF) land in ``document_texts`` with a ``ref`` like
  ``paragraph 12`` or ``page 3``
"""

from __future__ import annotations

import io
import logging
import xml.etree.ElementTree as ET
import zipfile
from dataclasses import dataclass, field
from pathlib import Path

import duckdb
import pandas as pd

from ..models import DocumentInfo
from .normalize import coerce_columns, decode_bytes, dedupe_identifiers, sanitize_identifier

logger = logging.getLogger(__name__)

SKIP_SUFFIXES = {".dtd"}
SKIP_NAMES = {".ds_store"}


@dataclass
class IngestResult:
    db_path: Path
    documents: list[DocumentInfo] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


def ingest_zip(zip_path: Path, work_dir: Path) -> IngestResult:
    """Extract ``zip_path`` into ``work_dir/extracted`` and load everything
    into ``work_dir/dossier.duckdb``."""
    logger.info("ingest_zip: extracting %s → %s", zip_path.name, work_dir)
    extract_dir = work_dir / "extracted"
    extract_dir.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(zip_path) as zf:
        root = extract_dir.resolve()
        for member in zf.infolist():
            destination = (extract_dir / member.filename).resolve()
            if not destination.is_relative_to(root):
                raise ValueError(f"unsafe path in ZIP archive: {member.filename}")
        zf.extractall(extract_dir)

    db_path = work_dir / "dossier.duckdb"
    if db_path.exists():
        db_path.unlink()
    con = duckdb.connect(str(db_path))
    result = IngestResult(db_path=db_path)

    con.execute(
        """
        CREATE TABLE documents (
            document_id VARCHAR, file VARCHAR, kind VARCHAR,
            table_name VARCHAR, row_count BIGINT
        )
        """
    )
    con.execute(
        """
        CREATE TABLE document_texts (
            document_id VARCHAR, file VARCHAR, ref VARCHAR, text VARCHAR
        )
        """
    )

    files = sorted(
        p for p in extract_dir.rglob("*")
        if p.is_file()
        and p.suffix.lower() not in SKIP_SUFFIXES
        and p.name.lower() not in SKIP_NAMES
        and "__macosx" not in str(p).lower()
    )
    logger.info("ingest_zip: %d files to load", len(files))

    # GDPdU index.xml files describe sibling txt tables; collect their schemas
    # first so the txt files are loaded with proper column names.
    gdpdu_schemas: dict[Path, tuple[str, list[str]]] = {}
    for index_file in (p for p in files if p.name.lower() == "index.xml"):
        try:
            gdpdu_schemas.update(_parse_gdpdu_index(index_file))
        except Exception as exc:  # noqa: BLE001 - ingestion must not die on one bad file
            result.warnings.append(f"failed to parse {index_file.name}: {exc}")

    doc_counter = 0
    for path in files:
        if path.name.lower() == "index.xml":
            continue
        rel = str(path.relative_to(extract_dir))
        doc_counter += 1
        document_id = f"doc-{doc_counter:03d}"
        try:
            infos = _load_file(con, path, rel, document_id, gdpdu_schemas)
            for info in infos:
                logger.info(
                    "  loaded %s → %s kind=%s rows=%s",
                    document_id,
                    rel,
                    info.kind,
                    info.row_count,
                )
        except Exception as exc:  # noqa: BLE001
            result.warnings.append(f"failed to ingest {rel}: {exc}")
            logger.exception("ingestion failed for %s", rel)
            infos = [DocumentInfo(document_id=document_id, file=rel, kind="other")]
        for info in infos:
            con.execute(
                "INSERT INTO documents VALUES (?, ?, ?, ?, ?)",
                [info.document_id, info.file, info.kind, info.table, info.row_count],
            )
            result.documents.append(info)

    con.close()
    logger.info(
        "ingest_zip complete — %d documents, %d warnings",
        len(result.documents),
        len(result.warnings),
    )
    return result


# ---------------------------------------------------------------------------
# Loaders
# ---------------------------------------------------------------------------


def _load_file(
    con: duckdb.DuckDBPyConnection,
    path: Path,
    rel: str,
    document_id: str,
    gdpdu_schemas: dict[Path, tuple[str, list[str]]],
) -> list[DocumentInfo]:
    suffix = path.suffix.lower()
    if path in gdpdu_schemas:
        return [_load_gdpdu_table(con, path, rel, document_id, *gdpdu_schemas[path])]
    if suffix == ".csv":
        return [_load_csv(con, path, rel, document_id)]
    if suffix == ".xlsx":
        return _load_xlsx(con, path, rel, document_id)
    if suffix == ".docx":
        return [_load_docx(con, path, rel, document_id)]
    if suffix == ".pdf":
        return [_load_pdf(con, path, rel, document_id)]
    if suffix == ".txt":
        # txt without a GDPdU schema: try semicolon CSV without header
        return [_load_csv(con, path, rel, document_id, header=None)]
    return [DocumentInfo(document_id=document_id, file=rel, kind="other")]


def _parse_gdpdu_index(index_file: Path) -> dict[Path, tuple[str, list[str]]]:
    """Return {data_file_path: (table_name, [column, ...])} for one index.xml."""
    tree = ET.parse(index_file)
    schemas: dict[Path, tuple[str, list[str]]] = {}
    for table in tree.getroot().iter("Table"):
        url = table.findtext("URL")
        name = table.findtext("Name") or (url or "table")
        if not url:
            continue
        columns = [
            col.findtext("Name") or f"col_{i}"
            for i, col in enumerate(table.iter("VariableColumn"))
        ]
        data_path = index_file.parent / url
        folder = sanitize_identifier(index_file.parent.name)
        table_name = f"{folder}__{sanitize_identifier(name)}"
        schemas[data_path] = (table_name, columns)
    return schemas


def _register_dataframe(
    con: duckdb.DuckDBPyConnection,
    df: pd.DataFrame,
    table_name: str,
    rel: str,
    document_id: str,
    kind: str,
) -> DocumentInfo:
    df = coerce_columns(df)
    con.register("_incoming", df)
    con.execute(f'CREATE TABLE "{table_name}" AS SELECT * FROM _incoming')
    con.unregister("_incoming")
    return DocumentInfo(
        document_id=document_id,
        file=rel,
        kind=kind,  # type: ignore[arg-type]
        table=table_name,
        row_count=len(df),
    )


def _read_delimited(path: Path, header: int | None, names: list[str] | None) -> pd.DataFrame:
    text = decode_bytes(path.read_bytes())
    df = pd.read_csv(
        io.StringIO(text),
        sep=";",
        header=header,
        names=names,
        dtype=str,
        keep_default_na=False,
        quotechar='"',
        engine="python",
    )
    df.columns = dedupe_identifiers([sanitize_identifier(str(c)) for c in df.columns])
    # _row_id = physical line number in the source file
    first_data_line = 1 if header is None else header + 2
    df.insert(0, "_row_id", range(first_data_line, first_data_line + len(df)))
    return df


def _load_gdpdu_table(
    con: duckdb.DuckDBPyConnection,
    path: Path,
    rel: str,
    document_id: str,
    table_name: str,
    columns: list[str],
) -> DocumentInfo:
    df = _read_delimited(path, header=None, names=columns)
    return _register_dataframe(con, df, table_name, rel, document_id, "gdpdu_table")


def _load_csv(
    con: duckdb.DuckDBPyConnection,
    path: Path,
    rel: str,
    document_id: str,
    header: int | None = 0,
) -> DocumentInfo:
    df = _read_delimited(path, header=header, names=None)
    table_name = sanitize_identifier(path.stem)
    return _register_dataframe(con, df, table_name, rel, document_id, "csv")


def _load_xlsx(
    con: duckdb.DuckDBPyConnection, path: Path, rel: str, document_id: str
) -> list[DocumentInfo]:
    sheets = pd.read_excel(path, sheet_name=None, dtype=str, keep_default_na=False)
    infos: list[DocumentInfo] = []
    for sheet_name, df in sheets.items():
        df.columns = dedupe_identifiers([sanitize_identifier(str(c)) for c in df.columns])
        # header is spreadsheet row 1, so data rows start at 2
        df.insert(0, "_row_id", range(2, 2 + len(df)))
        table_name = f"{sanitize_identifier(path.stem)}__{sanitize_identifier(sheet_name)}"
        info = _register_dataframe(con, df, table_name, rel, document_id, "xlsx_sheet")
        infos.append(info)
    return infos or [DocumentInfo(document_id=document_id, file=rel, kind="other")]


def _load_docx(
    con: duckdb.DuckDBPyConnection, path: Path, rel: str, document_id: str
) -> DocumentInfo:
    import docx  # lazy import

    document = docx.Document(str(path))
    count = 0
    for i, para in enumerate(document.paragraphs, start=1):
        text = para.text.strip()
        if not text:
            continue
        con.execute(
            "INSERT INTO document_texts VALUES (?, ?, ?, ?)",
            [document_id, rel, f"paragraph {i}", text],
        )
        count += 1
    for t_idx, table in enumerate(document.tables, start=1):
        for r_idx, row in enumerate(table.rows, start=1):
            cells = " | ".join(c.text.strip() for c in row.cells)
            if not cells.strip(" |"):
                continue
            con.execute(
                "INSERT INTO document_texts VALUES (?, ?, ?, ?)",
                [document_id, rel, f"table {t_idx} row {r_idx}", cells],
            )
            count += 1
    return DocumentInfo(document_id=document_id, file=rel, kind="docx", row_count=count)


def _load_pdf(
    con: duckdb.DuckDBPyConnection, path: Path, rel: str, document_id: str
) -> DocumentInfo:
    from pypdf import PdfReader  # lazy import

    reader = PdfReader(str(path))
    count = 0
    for page_no, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        con.execute(
            "INSERT INTO document_texts VALUES (?, ?, ?, ?)",
            [document_id, rel, f"page {page_no}", text],
        )
        count += 1
    return DocumentInfo(document_id=document_id, file=rel, kind="pdf", row_count=count)


# ---------------------------------------------------------------------------
# Schema overview for the agent prompt
# ---------------------------------------------------------------------------


def schema_overview(db_path: Path) -> str:
    """Compact plain-text description of every table for the system prompt."""
    con = duckdb.connect(str(db_path), read_only=True)
    try:
        docs = con.execute(
            "SELECT document_id, file, kind, table_name, row_count FROM documents ORDER BY document_id"
        ).fetchall()
        lines: list[str] = []
        for document_id, file, kind, table_name, row_count in docs:
            if table_name:
                cols = con.execute(
                    "SELECT column_name, data_type FROM information_schema.columns "
                    "WHERE table_name = ? ORDER BY ordinal_position",
                    [table_name],
                ).fetchall()
                col_desc = ", ".join(f"{c} {t}" for c, t in cols)
                lines.append(
                    f"- {table_name} ({row_count} rows) [{document_id}, file: {file}]: {col_desc}"
                )
            elif kind in ("docx", "pdf"):
                lines.append(
                    f"- document_texts entries for {document_id} (file: {file}, {kind}, {row_count} passages)"
                )
            else:
                lines.append(f"- {document_id} (file: {file}) not ingested as a table")
        return "\n".join(lines)
    finally:
        con.close()
